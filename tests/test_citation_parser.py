"""Tests for citation_parser.py."""
import zipfile
import shutil
import tempfile
import os
import pytest
from pathlib import Path
from docx import Document as DocxDocument


FIXTURES_DIR = Path(__file__).parent / "fixtures"


# ---------------------------------------------------------------------------
# _parse_paragraph — inline parsing
# ---------------------------------------------------------------------------

class TestParseParagraph:

    def test_parse_single_author_year(self, parser):
        results = parser._parse_paragraph("See (Smith, 2020).", 0, in_footnote=False)
        assert len(results) == 1
        assert results[0].authors == ["Smith"]
        assert results[0].year == "2020"

    def test_parse_two_citations(self, parser):
        results = parser._parse_paragraph(
            "As argued by (Jones, 2018) and (Brown, 2019).", 0, in_footnote=False
        )
        assert len(results) == 2

    def test_parse_multi_author(self, parser):
        results = parser._parse_paragraph(
            "(Garcia & Lopez, 2021)", 0, in_footnote=False
        )
        assert len(results) == 1
        assert "Garcia" in results[0].authors
        assert "Lopez" in results[0].authors

    def test_parse_et_al(self, parser):
        results = parser._parse_paragraph("(Smith et al., 2020)", 0, in_footnote=False)
        assert len(results) == 1
        assert results[0].authors == ["Smith"]

    @pytest.mark.parametrize("text", [
        "(Smith m.fl., 2020)",
        "(Smith m. fl., 2020)",
        "(Smith mfl., 2020)",
        "(Smith mfl, 2020)",
        "(Smith M.Fl., 2020)",
    ])
    def test_parse_mfl_variants(self, parser, text):
        """Swedish 'm.fl.' ('med flera') should be treated like 'et al.'"""
        results = parser._parse_paragraph(text, 0, in_footnote=False)
        assert len(results) == 1, f"Failed to parse {text!r}"
        assert results[0].authors == ["Smith"]
        assert results[0].year == "2020"
        assert results[0].has_et_al is True

    def test_parse_mfl_in_compound_citation(self, parser):
        """m.fl. in one unit should not confuse the other unit in a compound."""
        results = parser._parse_paragraph(
            "(Smith m.fl., 2020; Jones, 2021)", 0, in_footnote=False
        )
        assert len(results) == 2
        assert results[0].authors == ["Smith"] and results[0].has_et_al is True
        assert results[1].authors == ["Jones"] and results[1].has_et_al is False

    def test_mfl_inside_word_does_not_false_match(self, parser):
        """Strings like 'mflSthlm' inside a word must not be taken as m.fl."""
        results = parser._parse_paragraph(
            "mflSthlm är en förkortning (Smith, 2020).", 0, in_footnote=False
        )
        assert len(results) == 1
        assert results[0].has_et_al is False

    def test_parse_with_locator(self, parser):
        results = parser._parse_paragraph("(Smith, 2020:45)", 0, in_footnote=False)
        assert len(results) == 1
        assert results[0].year == "2020"
        assert results[0].locator == "45"

    def test_parse_with_cf_prefix(self, parser):
        results = parser._parse_paragraph("(cf., Smith, 2020)", 0, in_footnote=False)
        assert len(results) == 1
        assert results[0].authors == ["Smith"]

    def test_parse_diacritics(self, parser):
        results = parser._parse_paragraph("(Bååth, 2022)", 0, in_footnote=False)
        assert len(results) == 1
        assert results[0].authors == ["Bååth"]

    def test_no_false_positive_year_alone(self, parser):
        """A bare year without an author should not be a citation."""
        results = parser._parse_paragraph("This happened in 2020.", 0, in_footnote=False)
        assert len(results) == 0

    def test_in_footnote_flag(self, parser):
        results = parser._parse_paragraph("(Smith, 2020)", -1, in_footnote=True)
        assert results[0].in_footnote is True

    def test_semicolon_separated(self, parser):
        results = parser._parse_paragraph(
            "(Smith, 2020; Jones, 2019)", 0, in_footnote=False
        )
        assert len(results) == 2


# ---------------------------------------------------------------------------
# Regression: prefix tokens (see, cf., e.g., also) must not be parsed as
# authors. Cases collected from a real manuscript where the parser produced
# false negatives like "(f. Stone et al., 2018)" or "(see & Korsmeyer, 1999)".
# ---------------------------------------------------------------------------

class TestPrefixTokenRegressions:

    def test_cf_with_space_no_comma(self, parser):
        """`(cf. Stone et al. 2018)` — space, no comma. Bug: 'f. Stone' captured."""
        results = parser._parse_paragraph(
            "(cf. Stone et al. 2018)", 0, in_footnote=False
        )
        assert len(results) == 1
        assert results[0].authors == ["Stone"]
        assert results[0].year == "2018"
        assert results[0].has_et_al is True

    def test_cf_with_period_comma_space(self, parser):
        results = parser._parse_paragraph(
            "(cf., Quan and Wang 2004)", 0, in_footnote=False
        )
        assert len(results) == 1
        assert results[0].authors == ["Quan", "Wang"]
        assert results[0].year == "2004"

    def test_see_comma_single_author(self, parser):
        """`(see, Korsmeyer 1999)` — bug: 'see, Korsmeyer' as 2-author."""
        results = parser._parse_paragraph(
            "(see, Korsmeyer 1999)", 0, in_footnote=False
        )
        assert len(results) == 1
        assert results[0].authors == ["Korsmeyer"]
        assert results[0].year == "1999"

    def test_see_also_compound(self, parser):
        """`(see, also, Lane 2014; Leschziner 2007)` — bug: 'also & Lane' captured."""
        results = parser._parse_paragraph(
            "(see, also, Lane 2014; Leschziner 2007)", 0, in_footnote=False
        )
        assert len(results) == 2
        assert results[0].authors == ["Lane"]
        assert results[1].authors == ["Leschziner"]

    def test_eg_prefix(self, parser):
        results = parser._parse_paragraph(
            "(cf., e.g. Myhrvold 2011)", 0, in_footnote=False
        )
        assert len(results) == 1
        assert results[0].authors == ["Myhrvold"]
        assert results[0].year == "2011"

    def test_three_author_chicago_style(self, parser):
        """`(see, Cousins, O'Gorman, and Stierand 2010)` — bug: only Stierand survived."""
        results = parser._parse_paragraph(
            "(see, Cousins, O’Gorman, and Stierand 2010)", 0, in_footnote=False
        )
        assert len(results) == 1
        assert results[0].authors == ["Cousins", "O’Gorman", "Stierand"]
        assert results[0].year == "2010"

    def test_curly_apostrophe_in_name(self, parser):
        """O'Gorman with curly apostrophe (Word auto-correct) must parse."""
        results = parser._parse_paragraph(
            "(O’Gorman 2015)", 0, in_footnote=False
        )
        assert len(results) == 1
        assert results[0].authors == ["O’Gorman"]

    def test_hyphenated_unicode_name(self, parser):
        """`(Vargas-Sánchez and López-Guzmán 2022)` — both authors must survive."""
        results = parser._parse_paragraph(
            "(Vargas-Sánchez and López-Guzmán 2022)", 0, in_footnote=False
        )
        assert len(results) == 1
        assert results[0].authors == ["Vargas-Sánchez", "López-Guzmán"]

    def test_prose_paren_not_treated_as_citation(self, parser):
        """A discursive paren must not produce 'common humanity et al., 2006'."""
        results = parser._parse_paragraph(
            "(cf., also the role of a common humanity, rooted in nature, "
            "for the foundations of justifications in Boltanski and Thévenot 2006, 141–2)",
            0, in_footnote=False,
        )
        # Should find ONE citation: Boltanski & Thévenot 2006.
        assert len(results) == 1
        assert results[0].authors == ["Boltanski", "Thévenot"]
        assert results[0].year == "2006"

    def test_inline_and_between_authors(self, parser):
        """`Koponen and Niva (2020)` — bug: only 'Niva' captured because 'And' was a stopword.
        After fix, authors must be split into individual names so the matcher
        can compare each surname against Zotero entries separately."""
        results = parser._parse_paragraph(
            "As shown by Koponen and Niva (2020), artistic novelty…",
            0, in_footnote=False,
        )
        assert len(results) == 1
        assert results[0].authors == ["Koponen", "Niva"]
        assert results[0].year == "2020"

    def test_inline_and_at_sentence_start_still_blocked(self, parser):
        """`And Smith (2020)` at sentence start should NOT include 'And' as author."""
        results = parser._parse_paragraph(
            "First sentence. And Smith (2020) showed…",
            0, in_footnote=False,
        )
        assert len(results) == 1
        # "And" must not be part of the author name
        assert results[0].authors == ["Smith"]

    def test_possessive_with_curly_apostrophe(self, parser):
        """`Pellegrino's (2021)` with curly apostrophe must drop the 's."""
        results = parser._parse_paragraph(
            "Pellegrino’s (2021) attempt to justify…",
            0, in_footnote=False,
        )
        assert len(results) == 1
        assert results[0].authors == ["Pellegrino"]

    def test_initial_without_period(self, parser):
        """`M Nilsson (2020)` — initial without period must be stripped from
        the surname so matching against 'Nilsson' in Zotero succeeds."""
        results = parser._parse_paragraph(
            "an interesting nuance offered by M Nilsson (2020, 67–76).",
            0, in_footnote=False,
        )
        assert len(results) == 1
        # Author may be captured verbatim ("M Nilsson") but matcher key
        # must reduce to just the surname.
        from citation_parser import _first_surname_key
        assert _first_surname_key(results[0].authors[0]) == "nilsson"

    def test_narrative_with_locator_and_co_citation(self, parser):
        """`Lane (2014, 20; cf. Bildtgård 2010)` — author lives outside the
        paren; the leading year + page must still produce a Lane citation
        alongside the Bildtgård sub-citation."""
        results = parser._parse_paragraph(
            "I follow Lane (2014, 20; cf. Bildtgård 2010) in arguing.",
            0, in_footnote=False,
        )
        displays = [r.display_text for r in results]
        assert any(d == "(Lane, 2014)" for d in displays), displays
        assert any(d == "(Bildtgård, 2010)" for d in displays), displays
        lane = next(r for r in results if r.display_text == "(Lane, 2014)")
        assert lane.locator == "20"

    def test_real_compound_citation_intact(self, parser):
        """`(Batat 2021; Madeira et al. 2020; Vargas-Sánchez and López-Guzmán 2022)`
        all three must parse."""
        results = parser._parse_paragraph(
            "(Batat 2021; Madeira et al. 2020; Vargas-Sánchez and López-Guzmán 2022)",
            0, in_footnote=False,
        )
        assert len(results) == 3
        assert results[0].authors == ["Batat"] and results[0].year == "2021"
        assert results[1].authors == ["Madeira"] and results[1].has_et_al
        assert results[2].authors == ["Vargas-Sánchez", "López-Guzmán"]


# ---------------------------------------------------------------------------
# Zotero / EndNote field code stripping
# ---------------------------------------------------------------------------

def _make_docx_with_zotero_field(text_inside_field: str) -> str:
    """Create a .docx with one paragraph containing a Zotero ADDIN field code
    whose visible result text is text_inside_field. Returns temp file path."""
    from lxml import etree
    from docx.oxml.ns import qn

    doc = DocxDocument()
    doc.add_paragraph("")  # ensure at least one paragraph exists
    p = doc.paragraphs[0]._element

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def w(tag):
        return f"{{{W}}}{tag}"

    def make_run(text):
        r = etree.SubElement(p, w("r"))
        t = etree.SubElement(r, w("t"))
        t.text = text
        return r

    def make_fldchar(typ):
        r = etree.SubElement(p, w("r"))
        fc = etree.SubElement(r, w("fldChar"))
        fc.set(w("fldCharType"), typ)
        return r

    def make_instrtext(instr):
        r = etree.SubElement(p, w("r"))
        it = etree.SubElement(r, w("instrText"))
        it.text = instr
        return r

    make_run("Before ")
    make_fldchar("begin")
    make_instrtext(' ADDIN ZOTERO_ITEM CSL_CITATION {"citationID":"test"} ')
    make_fldchar("separate")
    make_run(text_inside_field)
    make_fldchar("end")
    make_run(" after.")

    tmp = tempfile.mktemp(suffix=".docx")
    doc.save(tmp)
    return tmp


class TestFieldStripping:

    def test_zotero_field_stripped_and_visible(self):
        """After stripping, the visible citation text becomes plain text."""
        from citation_parser import CitationParser
        tmp = _make_docx_with_zotero_field("(Smith, 2020)")
        try:
            cp = CitationParser(tmp)
            cits = cp.parse()
            assert any(
                "Smith" in c.authors for c in cits
            ), f"Expected Smith citation, got: {cits}"
        finally:
            os.unlink(tmp)

    def test_original_not_modified(self):
        """Parsing must not modify the original file."""
        from citation_parser import CitationParser
        tmp = _make_docx_with_zotero_field("(Jones, 2020)")
        try:
            mtime_before = os.path.getmtime(tmp)
            CitationParser(tmp).parse()
            mtime_after = os.path.getmtime(tmp)
            assert mtime_before == mtime_after
        finally:
            os.unlink(tmp)


# ---------------------------------------------------------------------------
# Tracked changes — regression for the duplicate-citation bug
# ---------------------------------------------------------------------------

def _make_docx_with_tracked_insertion(citation_text: str) -> str:
    """Create a .docx where a Zotero field code is wrapped in a <w:ins> element
    (tracked insertion). This was the root cause of the duplicate-citation bug."""
    from lxml import etree

    doc = DocxDocument()
    doc.add_paragraph("")  # ensure at least one paragraph exists
    p = doc.paragraphs[0]._element

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def w(tag):
        return f"{{{W}}}{tag}"

    # Wrap field code runs inside <w:ins>
    ins = etree.SubElement(p, w("ins"))
    ins.set(w("id"), "1")
    ins.set(w("author"), "Test")
    ins.set(w("date"), "2024-01-01T00:00:00Z")

    def make_run_in(parent, text):
        r = etree.SubElement(parent, w("r"))
        t = etree.SubElement(r, w("t"))
        t.text = text
        return r

    def make_fldchar_in(parent, typ):
        r = etree.SubElement(parent, w("r"))
        fc = etree.SubElement(r, w("fldChar"))
        fc.set(w("fldCharType"), typ)

    def make_instrtext_in(parent, instr):
        r = etree.SubElement(parent, w("r"))
        it = etree.SubElement(r, w("instrText"))
        it.text = instr

    make_fldchar_in(ins, "begin")
    make_instrtext_in(ins, ' ADDIN ZOTERO_ITEM CSL_CITATION {"citationID":"t1"} ')
    make_fldchar_in(ins, "separate")
    make_run_in(ins, citation_text)
    make_fldchar_in(ins, "end")

    # Also add plain text after
    r2 = etree.SubElement(p, w("r"))
    t2 = etree.SubElement(r2, w("t"))
    t2.text = " Some surrounding text."

    tmp = tempfile.mktemp(suffix=".docx")
    doc.save(tmp)
    return tmp


class TestTrackedChanges:

    def test_no_duplicate_from_tracked_insertion(self):
        """Field code inside <w:ins> must be stripped and not produce duplicate citations."""
        from citation_parser import CitationParser
        tmp = _make_docx_with_tracked_insertion("(Smith, 2020)")
        try:
            cits = CitationParser(tmp).parse()
            smith_cits = [c for c in cits if "Smith" in c.authors]
            assert len(smith_cits) == 1, (
                f"Expected exactly 1 Smith citation, got {len(smith_cits)}: {smith_cits}"
            )
        finally:
            os.unlink(tmp)

    def test_tracked_changes_accepted_before_strip(self):
        """The citation inside <w:ins> must be found (not silently ignored)."""
        from citation_parser import CitationParser
        tmp = _make_docx_with_tracked_insertion("(Brown, 2019)")
        try:
            cits = CitationParser(tmp).parse()
            assert any("Brown" in c.authors for c in cits), (
                "Citation inside <w:ins> was not parsed"
            )
        finally:
            os.unlink(tmp)
