"""Tests for field_writer.py."""
import os
import tempfile
import pytest
from pathlib import Path
from docx import Document as DocxDocument


FIXTURES_DIR = Path(__file__).parent / "fixtures"
TEST_DB = FIXTURES_DIR / "test_library.sqlite"


def _make_simple_docx(text: str) -> str:
    """Create a minimal .docx with one paragraph of text. Returns temp path."""
    doc = DocxDocument()
    doc.add_paragraph(text)
    tmp = tempfile.mktemp(suffix=".docx")
    doc.save(tmp)
    return tmp


def _load_library():
    from zotero_client import ZoteroClient
    return ZoteroClient(str(TEST_DB)).load_library()


def _match(library, text):
    from citation_parser import CitationParser
    from matcher import CitationMatcher
    tmp = _make_simple_docx(text)
    try:
        cits = CitationParser(tmp).parse()
        results = CitationMatcher(library).match_all(cits)
        return tmp, results
    except Exception:
        os.unlink(tmp)
        raise


class TestOutputFile:

    def test_original_not_overwritten(self):
        """Output must be saved to a new file; original unchanged."""
        from field_writer import write_zotero_document
        library = _load_library()
        src, results = _match(library, "See (Smith, 2020) for details.")
        mtime_before = os.path.getmtime(src)
        out = src.replace(".docx", "_zotero.docx")
        try:
            write_zotero_document(src, out, results)
            assert os.path.exists(out), "Output file was not created"
            assert os.path.getmtime(src) == mtime_before, "Original file was modified"
        finally:
            os.unlink(src)
            if os.path.exists(out):
                os.unlink(out)

    def test_output_contains_field_code(self):
        """Output .docx must contain a Zotero ADDIN field code for matched citation."""
        from field_writer import write_zotero_document
        library = _load_library()
        src, results = _match(library, "See (Smith, 2020) for details.")
        out = src.replace(".docx", "_zotero.docx")
        try:
            write_zotero_document(src, out, results)
            W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            doc = DocxDocument(out)
            found = any(
                elem.text and "ZOTERO_ITEM" in elem.text
                for elem in doc.element.body.iter(f"{{{W}}}instrText")
            )
            assert found, "No Zotero field code found in output document"
        finally:
            os.unlink(src)
            if os.path.exists(out):
                os.unlink(out)

    def test_unmatched_citation_stays_plain(self):
        """Unmatched citations must remain as plain text in the output."""
        from field_writer import write_zotero_document
        library = _load_library()
        src, results = _match(library, "See (Zzzyxq, 9999) for details.")
        out = src.replace(".docx", "_zotero.docx")
        try:
            write_zotero_document(src, out, results)
            doc = DocxDocument(out)
            full_text = " ".join(p.text for p in doc.paragraphs)
            assert "Zzzyxq" in full_text, "Unmatched citation text was removed"
        finally:
            os.unlink(src)
            if os.path.exists(out):
                os.unlink(out)


class TestSharedSpanRegression:
    """Regression tests for github issue #1: text adjacent to a compound
    citation could be silently truncated when two sub-citations resolved to
    the same (paragraph, char_start, char_end) span and were then processed
    as two separate calls to ``_replace_citation_in_paragraph``."""

    @staticmethod
    def _make_fake_item(cit, key):
        from zotero_client import ZoteroItem
        return ZoteroItem(
            key=key, item_type_zotero='journalArticle',
            item_type_csl='article-journal',
            title=f'Fake {(cit.authors or ["x"])[0]}', year=cit.year,
            authors=cit.authors or ['Anon'],
            library_id=1, user_id='0', item_id=99000,
            csl_data={'id': key, 'type': 'article-journal'},
        )

    def test_compound_citation_preserves_surrounding_text(self):
        """End-to-end: a compound citation `(A, 2020; B, 2021)` in a paragraph
        with text both before and after must not damage the surrounding text.
        This exercises the normal grouping path."""
        from citation_parser import CitationParser
        from matcher import MatchResult
        from field_writer import write_zotero_document

        text = (
            "Lorem ipsum dolor sit amet, consectetur adipiscing elit. "
            "(Smith, 2020; Jones, 2021) Sed do eiusmod tempor incididunt ut "
            "labore et dolore magna aliqua."
        )
        src = _make_simple_docx(text)
        out = src.replace(".docx", "_zotero.docx")
        try:
            cits = CitationParser(src).parse()
            assert len(cits) >= 2, "Expected compound citation to parse as 2 units"
            results = [
                MatchResult(
                    citation=c,
                    zotero_item=self._make_fake_item(c, f"K{i}"),
                    confidence=0.95, matched=True,
                )
                for i, c in enumerate(cits)
            ]
            write_zotero_document(src, out, results)
            full_text = DocxDocument(out).paragraphs[0].text
            assert "Lorem ipsum dolor sit amet, consectetur adipiscing elit." in full_text
            assert "Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua." in full_text
        finally:
            os.unlink(src)
            if os.path.exists(out):
                os.unlink(out)

    def test_duplicate_span_calls_do_not_corrupt_text(self):
        """Direct test of the corruption mechanism: call
        ``_replace_citation_in_paragraph`` twice with the same span.

        The first call inserts a field; the second must detect that the
        target span is already inside a Zotero field region and return
        without touching the surrounding prefix/suffix runs.
        """
        from citation_parser import CitationParser
        from field_writer import FieldWriter
        from matcher import MatchResult

        text = (
            "Before text that must be preserved. "
            "(Smith, 2020; Jones, 2021) "
            "After text that must also be preserved."
        )
        src = _make_simple_docx(text)
        try:
            cits = CitationParser(src).parse()
            shared = [c for c in cits
                      if c.char_start == cits[0].char_start
                      and c.char_end == cits[0].char_end]
            assert len(shared) >= 2, "Expected ≥2 citations with shared span"

            m1 = MatchResult(citation=shared[0],
                             zotero_item=self._make_fake_item(shared[0], "K1"),
                             confidence=0.95, matched=True)
            m2 = MatchResult(citation=shared[1],
                             zotero_item=self._make_fake_item(shared[1], "K2"),
                             confidence=0.95, matched=True)

            writer = FieldWriter(src)
            para = writer._all_paras[shared[0].paragraph_idx]
            writer._replace_citation_in_paragraph(para, m1.citation, [m1])
            writer._replace_citation_in_paragraph(para, m2.citation, [m2])
            out = src.replace(".docx", "_dup.docx")
            writer.save(out)

            full_text = DocxDocument(out).paragraphs[0].text
            assert "Before text that must be preserved." in full_text, \
                f"Prefix was damaged: {full_text!r}"
            assert "After text that must also be preserved." in full_text, \
                f"Suffix was damaged: {full_text!r}"
            if os.path.exists(out):
                os.unlink(out)
        finally:
            os.unlink(src)

    def test_vietnamese_compound_citation_issue_1(self):
        """Exact reproduction of github issue #1: a Vietnamese paragraph
        with a compound citation. The Vietnamese text between the citation
        and the end of the paragraph must be preserved verbatim."""
        from citation_parser import CitationParser
        from matcher import MatchResult
        from field_writer import write_zotero_document

        text = (
            "Tổng quan các nghiên cứu trong và ngoài nước cho thấy quốc tế "
            "hóa giáo dục y khoa và các chương trình trao đổi quốc tế đang "
            "phát triển mạnh mẽ, đóng vai trò quan trọng trong nâng cao năng "
            "lực chuyên môn và hội nhập toàn cầu của người học. "
            "(Nguyen, 2024; Alshardan & Sabbagh, 2014) "
            "Tuy nhiên, vẫn còn tồn tại nhiều khoảng trống nghiên cứu cần "
            "được làm rõ."
        )
        src = _make_simple_docx(text)
        out = src.replace(".docx", "_zotero.docx")
        try:
            cits = CitationParser(src).parse()
            results = [
                MatchResult(
                    citation=c,
                    zotero_item=self._make_fake_item(c, f"K{i}"),
                    confidence=0.95, matched=True,
                )
                for i, c in enumerate(cits)
            ]
            write_zotero_document(src, out, results)
            full_text = DocxDocument(out).paragraphs[0].text
            # The suffix the user reported was being truncated:
            assert "Tuy nhiên, vẫn còn tồn tại nhiều khoảng" in full_text
            assert "trống nghiên cứu cần được làm rõ." in full_text
            assert len(full_text) == len(text), \
                f"Paragraph length changed: {len(full_text)} vs {len(text)}"
        finally:
            os.unlink(src)
            if os.path.exists(out):
                os.unlink(out)
