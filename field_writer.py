"""
field_writer.py
Replaces plain-text citations in a .docx or .odt with Zotero field codes.
Operates directly on the Word XML (via lxml) or ODF XML (via odfpy).
"""
from __future__ import annotations

import copy
import json
import random
import re
import string
from collections import defaultdict
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from citation_parser import Citation
from matcher import MatchResult
from zotero_client import ZoteroItem

# ---------------------------------------------------------------------------
# Optional ODT support (odfpy)
# ---------------------------------------------------------------------------
try:
    from odf.opendocument import load as _odf_load
    from odf.text import P as _OdfP, ReferenceMarkStart, ReferenceMarkEnd
    from odf.element import Text as _OdfText
    _ODT_AVAILABLE = True
except ImportError:
    _ODT_AVAILABLE = False

_ODT_NOTE_QNAME = ('urn:oasis:names:tc:opendocument:xmlns:text:1.0', 'note')
_ODT_NOTE_BODY_QNAMES = {
    ('urn:oasis:names:tc:opendocument:xmlns:text:1.0', 'footnote-body'),
    ('urn:oasis:names:tc:opendocument:xmlns:text:1.0', 'endnote-body'),
}


# ---------------------------------------------------------------------------
# XML namespace helpers
# ---------------------------------------------------------------------------

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"


def _w(tag: str) -> str:
    return f"{W}{tag}"


def accept_tracked_changes(body) -> None:
    """Accept all tracked changes in a document body element.

    - <w:del> blocks are removed entirely (deleted text discarded).
    - <w:ins> wrappers are unwrapped (inserted content kept as direct children).
    - <w:rPrChange> / <w:pPrChange> elements are removed (new formatting kept).

    Call on ``doc.element.body`` before parsing or writing so that both
    CitationParser and FieldWriter see identical, clean XML.
    """
    W_INS        = _w("ins")
    W_DEL        = _w("del")
    W_RPR_CHANGE = _w("rPrChange")
    W_PPR_CHANGE = _w("pPrChange")

    # 1. Remove formatting-change markers (keep new formatting, discard old)
    for tag in (W_RPR_CHANGE, W_PPR_CHANGE):
        for elem in body.findall(f".//{tag}"):
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)

    # 2. Remove deleted content
    for elem in body.findall(f".//{W_DEL}"):
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)

    # 3. Unwrap insertions — process in reverse document order so that
    #    nested <w:ins> elements are handled correctly.
    for elem in reversed(body.findall(f".//{W_INS}")):
        parent = elem.getparent()
        if parent is None:
            continue
        siblings = list(parent)
        try:
            idx = siblings.index(elem)
        except ValueError:
            continue
        for i, child in enumerate(list(elem)):
            parent.insert(idx + i, child)
        parent.remove(elem)


def _make_run(parent_rpr=None) -> etree._Element:
    """Create a bare <w:r> element, optionally copying run properties."""
    r = etree.Element(_w("r"))
    if parent_rpr is not None:
        r.append(copy.deepcopy(parent_rpr))
    return r


def _random_citation_id(length: int = 8) -> str:
    return "".join(random.choices(string.ascii_letters + string.digits, k=length))


# ---------------------------------------------------------------------------
# Field code builder
# ---------------------------------------------------------------------------

def _build_csl_citation_json(
    matches: List[MatchResult], formatted_text: str
) -> dict:
    """Build a Zotero CSL_CITATION dict with one citationItem per match."""
    citation_items = []
    for m in matches:
        item: ZoteroItem = m.zotero_item  # type: ignore[assignment]
        item_dict = {
            "id": item.item_id,
            "uris": [item.uri],
            "itemData": item.csl_data,
        }
        if m.citation.locator:
            item_dict["locator"] = m.citation.locator
            item_dict["label"] = m.citation.locator_label
        if m.citation.prefix:
            item_dict["prefix"] = m.citation.prefix
        if m.citation.suffix:
            item_dict["suffix"] = m.citation.suffix
        citation_items.append(item_dict)
    return {
        "citationID": _random_citation_id(),
        "properties": {
            "formattedCitation": formatted_text,
            "plainCitation": formatted_text,
            "noteIndex": 0,
        },
        "citationItems": citation_items,
    }


def build_zotero_field_xml(
    matches: List[MatchResult],
    formatted_text: str,
    parent_rpr=None,
) -> List[etree._Element]:
    """
    Return a list of 5 <w:r> elements that form a Zotero ADDIN field:

        <w:r><w:fldChar begin/></w:r>
        <w:r><w:instrText> ADDIN ZOTERO_ITEM CSL_CITATION {...} </w:instrText></w:r>
        <w:r><w:fldChar separate/></w:r>
        <w:r><w:t>{formatted_text}</w:t></w:r>
        <w:r><w:fldChar end/></w:r>

    *matches* may contain one or more MatchResults.  When a compound
    citation block has multiple matched references they are all included
    in the same field code as separate citationItems.
    """
    csl_citation = _build_csl_citation_json(matches, formatted_text)
    instr_text = " ADDIN ZOTERO_ITEM CSL_CITATION " + json.dumps(csl_citation) + " "

    runs: List[etree._Element] = []

    # 1. fldChar begin
    r1 = _make_run(parent_rpr)
    fc1 = etree.SubElement(r1, _w("fldChar"))
    fc1.set(_w("fldCharType"), "begin")
    runs.append(r1)

    # 2. instrText
    r2 = _make_run(parent_rpr)
    it = etree.SubElement(r2, _w("instrText"))
    it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    it.text = instr_text
    runs.append(r2)

    # 3. fldChar separate
    r3 = _make_run(parent_rpr)
    fc3 = etree.SubElement(r3, _w("fldChar"))
    fc3.set(_w("fldCharType"), "separate")
    runs.append(r3)

    # 4. displayed text
    r4 = _make_run(parent_rpr)
    t = etree.SubElement(r4, _w("t"))
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = formatted_text
    runs.append(r4)

    # 5. fldChar end
    r5 = _make_run(parent_rpr)
    fc5 = etree.SubElement(r5, _w("fldChar"))
    fc5.set(_w("fldCharType"), "end")
    runs.append(r5)

    return runs


# ---------------------------------------------------------------------------
# Compound citation grouping
# ---------------------------------------------------------------------------

def _group_by_span(matched: List[MatchResult]) -> List[List[MatchResult]]:
    """Group matched results that share the same text span (paragraph_idx,
    char_start, char_end).  Each group becomes one Zotero field code.

    Returns a list of groups (each group is a list of MatchResults).
    """
    groups: Dict[Tuple[int, int, int], List[MatchResult]] = defaultdict(list)
    for r in matched:
        key = (r.citation.paragraph_idx, r.citation.char_start, r.citation.char_end)
        groups[key].append(r)
    return list(groups.values())


# ---------------------------------------------------------------------------
# Document writer
# ---------------------------------------------------------------------------

class FieldWriter:
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self._doc = Document(docx_path)
        # Accept all tracked changes so that <w:ins>-wrapped runs become direct
        # children of <w:p>.  Without this, citations inside tracked insertions
        # are detected by the parser but skipped during replacement because the
        # XML manipulation requires runs to be direct paragraph children.
        accept_tracked_changes(self._doc.element.body)
        # Strip existing Zotero/EndNote field codes, replacing each with its
        # visible text.  CitationParser does the same before computing char
        # offsets, so both sides see identical plain text and offsets align.
        # Must run before _all_paras is built, so operate on raw XML.
        for p_elem in self._doc.element.body.iter(_w("p")):
            self._strip_fields_in_element(p_elem)
        # Build the same expanded paragraph list as CitationParser:
        # body-level paragraphs + paragraphs inside table cells.
        self._all_paras = list(self._doc.paragraphs)
        for table in self._doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para not in self._all_paras:
                            self._all_paras.append(para)

    def _strip_zotero_fields(self) -> None:
        """Replace existing Zotero and EndNote field codes with their visible text.

        A citation manager field code in Word XML looks like:
          <w:r><w:fldChar w:fldCharType="begin"/></w:r>
          <w:r><w:instrText> ADDIN ZOTERO_ITEM CSL_CITATION {...} </w:instrText></w:r>
          <w:r><w:fldChar w:fldCharType="separate"/></w:r>
          <w:r><w:t>visible citation text</w:t></w:r>   ← may be multiple runs
          <w:r><w:fldChar w:fldCharType="end"/></w:r>

        EndNote uses the same structure with "ADDIN EN.CITE" (inline citations)
        and "ADDIN EN.REFLIST" (bibliography) in the instrText.

        This method finds each such sequence and replaces the entire thing with
        just the visible-text run(s), converting the citation back to plain text.
        """
        for para in self._all_paras:
            self._strip_fields_in_element(para._p)

    # Markers in instrText that identify citation-manager field codes
    _CITATION_FIELD_MARKERS = ("ZOTERO_ITEM", "EN.CITE", "EN.REFLIST")

    @staticmethod
    def _strip_fields_in_element(p_elem) -> None:
        """Strip Zotero and EndNote field codes from a single <w:p> element."""
        w_r = _w("r")
        w_fldChar = _w("fldChar")
        w_instrText = _w("instrText")
        w_fldCharType = _w("fldCharType")

        while True:
            children = list(p_elem)
            begin_idx = None
            is_citation_field = False
            sep_idx = None

            for i, child in enumerate(children):
                if child.tag != w_r:
                    continue
                fc = child.find(w_fldChar)
                if fc is not None:
                    ftype = fc.get(w_fldCharType)
                    if ftype == "begin":
                        begin_idx = i
                        is_citation_field = False
                        sep_idx = None
                    elif ftype == "separate" and begin_idx is not None:
                        sep_idx = i
                    elif ftype == "end" and begin_idx is not None:
                        if is_citation_field and sep_idx is not None:
                            # Found a complete citation-manager field code.
                            # Visible runs are between sep_idx+1 and i-1.
                            visible_runs = children[sep_idx + 1: i]
                            # Remove all field code elements (begin..end inclusive)
                            to_remove = children[begin_idx: i + 1]
                            insert_pos = begin_idx
                            for elem in to_remove:
                                p_elem.remove(elem)
                            # Re-insert only the visible text runs
                            for j, vr in enumerate(visible_runs):
                                p_elem.insert(insert_pos + j, vr)
                            break  # restart scan (indices changed)
                        begin_idx = None
                        sep_idx = None
                else:
                    # Check if this run has instrText from a citation manager
                    it = child.find(w_instrText)
                    if it is not None and begin_idx is not None:
                        if it.text and any(
                            m in it.text
                            for m in FieldWriter._CITATION_FIELD_MARKERS
                        ):
                            is_citation_field = True
            else:
                # Loop completed without break — no more Zotero fields
                break

    def apply_matches(
        self,
        results: List[MatchResult],
        accepted_suggestions: Optional[set] = None,
    ) -> None:
        """
        Replace each matched citation's plain text with a Zotero field code.
        Also replaces accepted suggestions (identified by their index in results).
        Modifies self._doc in place.
        Only processes body paragraphs (footnote support is a future enhancement).

        Compound citations (multiple references sharing the same parenthesized
        span) are grouped into a single field code with multiple citationItems.
        """
        accepted = accepted_suggestions or set()
        matched = [
            r for i, r in enumerate(results)
            if (r.matched or (r.is_suggestion and i in accepted))
            and not r.citation.in_footnote
        ]

        # Group by shared span so compound citations become one field code
        groups = _group_by_span(matched)

        # Sort groups by paragraph then reverse char position (process
        # later positions first so earlier offsets stay valid)
        groups.sort(key=lambda g: (g[0].citation.paragraph_idx,
                                   -g[0].citation.char_start))

        for group in groups:
            cit = group[0].citation          # representative citation
            para_idx = cit.paragraph_idx
            if para_idx < 0 or para_idx >= len(self._all_paras):
                continue
            para = self._all_paras[para_idx]
            self._replace_citation_in_paragraph(para, cit, group)

    def save(self, output_path: str) -> None:
        self._doc.save(output_path)

    # ------------------------------------------------------------------
    # Core replacement logic
    # ------------------------------------------------------------------

    def _replace_citation_in_paragraph(
        self, para, citation: Citation, group: List[MatchResult]
    ) -> None:
        """
        Find the span [char_start, char_end) in the paragraph's run sequence
        and replace it with Zotero field runs.

        *group* is a list of MatchResults that share the same span (compound
        citation).  They become a single field code with multiple citationItems.
        """
        # Build a flat list of (run_elem, char_offset_start, char_offset_end)
        run_spans = self._get_run_spans(para)
        if not run_spans:
            return

        start = citation.char_start
        end = citation.char_end
        formatted = citation.raw_text

        # Find which runs are covered by [start, end)
        first_run_idx = None
        last_run_idx = None

        for i, (run_elem, rs, re_, _in_ins) in enumerate(run_spans):
            if rs < end and re_ > start:
                if first_run_idx is None:
                    first_run_idx = i
                last_run_idx = i

        if first_run_idx is None:
            return

        # Get parent <w:rPr> from the first involved run (to preserve formatting)
        first_run_elem = run_spans[first_run_idx][0]
        parent_rpr = first_run_elem.find(_w("rPr"))

        # Build field XML runs (all items in the group share one field code)
        field_runs = build_zotero_field_xml(group, formatted, parent_rpr)

        # We'll rebuild the paragraph's run list.
        # Strategy: split first and last runs at the boundary, keep outer text,
        # replace the middle with field runs.

        para_xml = para._p  # the <w:p> element

        # Build prefix run (text before the citation in first_run)
        first_elem, first_rs, first_re, _ = run_spans[first_run_idx]
        prefix_text = self._get_run_text(first_elem)[: max(0, start - first_rs)]

        last_elem, last_rs, last_re, _ = run_spans[last_run_idx]
        suffix_text = self._get_run_text(last_elem)[max(0, end - last_rs) :]

        # Build prefix run element
        prefix_run = None
        if prefix_text:
            prefix_run = copy.deepcopy(first_elem)
            self._set_run_text(prefix_run, prefix_text)

        # Build suffix run element
        suffix_run = None
        if suffix_text:
            suffix_run = copy.deepcopy(last_elem)
            self._set_run_text(suffix_run, suffix_text)

        # Insert position: just before first_run in the paragraph
        insert_pos = list(para_xml).index(first_elem)

        # Remove the runs that are fully inside [first_run_idx, last_run_idx]
        for idx in range(first_run_idx, last_run_idx + 1):
            elem = run_spans[idx][0]
            if elem in para_xml:
                para_xml.remove(elem)

        # Re-insert: prefix + field runs + suffix
        new_elems: list = []
        if prefix_run is not None:
            new_elems.append(prefix_run)
        new_elems.extend(field_runs)
        if suffix_run is not None:
            new_elems.append(suffix_run)

        for j, elem in enumerate(new_elems):
            para_xml.insert(insert_pos + j, elem)

    # ------------------------------------------------------------------
    # Run text helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _get_run_spans(para) -> list:
        """
        Return list of (run_element, char_start, char_end, in_ins) tuples
        for every <w:r> in the paragraph, in document order.

        Mirrors the logic of citation_parser._get_paragraph_text so that
        char offsets are identical to those stored in Citation objects:
          - Includes <w:r> elements inside <w:ins> (tracked insertions)
          - Excludes <w:r> elements inside <w:del>
          - Excludes <w:r> elements inside existing Zotero field result regions
            (between fldChar separate and fldChar end of a ZOTERO_ITEM field)

        in_ins is True when the run is a child of <w:ins> rather than <w:p>.
        """
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        del_tag     = _w("del")
        ins_tag     = _w("ins")
        r_tag       = _w("r")
        fldChar_tag = _w("fldChar")
        instr_tag   = _w("instrText")

        spans: list = []
        offset = 0
        in_zotero_result = False
        zotero_field     = False

        for elem in para._p.iter():
            if elem.tag == fldChar_tag:
                ftype = elem.get(_w("fldCharType"), "")
                if ftype == "begin":
                    zotero_field = False
                    in_zotero_result = False
                elif ftype == "separate":
                    if zotero_field:
                        in_zotero_result = True
                elif ftype == "end":
                    in_zotero_result = False
                    zotero_field = False

            elif elem.tag == instr_tag:
                if elem.text and "ZOTERO_ITEM" in elem.text:
                    zotero_field = True

            elif elem.tag == r_tag:
                if in_zotero_result:
                    continue
                # Skip runs inside <w:del>
                ancestor = elem.getparent()
                in_del = False
                while ancestor is not None and ancestor != para._p:
                    if ancestor.tag == del_tag:
                        in_del = True
                        break
                    ancestor = ancestor.getparent()
                if in_del:
                    continue
                parent = elem.getparent()
                in_ins = parent is not None and parent.tag == ins_tag
                text = FieldWriter._get_run_text(elem)
                spans.append((elem, offset, offset + len(text), in_ins))
                offset += len(text)

        return spans

    @staticmethod
    def _get_run_text(run_elem: etree._Element) -> str:
        texts = []
        for t in run_elem.findall(_w("t")):
            texts.append(t.text or "")
        return "".join(texts)

    @staticmethod
    def _set_run_text(run_elem: etree._Element, text: str) -> None:
        for t in run_elem.findall(_w("t")):
            run_elem.remove(t)
        if text:
            t_elem = etree.SubElement(run_elem, _w("t"))
            t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t_elem.text = text


# ---------------------------------------------------------------------------
# ODT helpers
# ---------------------------------------------------------------------------

def _odt_text_nodes(elem, skip_qnames=None):
    """Yield (text_node, parent_elem) pairs in document order, skipping elements
    whose qname is in skip_qnames."""
    skip_qnames = skip_qnames or set()
    for child in elem.childNodes:
        if child.nodeType == 3:          # TEXT_NODE
            yield (child, elem)
        elif child.nodeType == 1:        # ELEMENT_NODE
            if hasattr(child, 'qname') and child.qname in skip_qnames:
                continue
            yield from _odt_text_nodes(child, skip_qnames)


def _inject_odt_marks(para, char_start: int, char_end: int, mark_name: str) -> bool:
    """
    Wrap text[char_start:char_end] in an ODF paragraph element with Zotero
    reference marks.  Returns True if the injection succeeded.
    Only handles the common case where the citation span falls entirely within
    a single text node (which covers nearly all unprocessed documents).

    Note: odfpy's removeChild/insertBefore reject Text nodes via internal cache
    assertions, so we manipulate childNodes directly.
    """
    text_nodes = list(_odt_text_nodes(para, skip_qnames={_ODT_NOTE_QNAME}))

    offset = 0
    for node, parent in text_nodes:
        node_len = len(node.data)
        node_end = offset + node_len

        if offset <= char_start and char_end <= node_end:
            rel_start = char_start - offset
            rel_end   = char_end   - offset

            before_text = node.data[:rel_start]
            cited_text  = node.data[rel_start:rel_end]
            after_text  = node.data[rel_end:]

            node_idx = parent.childNodes.index(node)

            mark_start = ReferenceMarkStart(**{'name': mark_name})
            mark_end   = ReferenceMarkEnd(**{'name': mark_name})

            new_nodes = []
            if before_text:
                new_nodes.append(_OdfText(before_text))
            new_nodes.append(mark_start)
            if cited_text:
                new_nodes.append(_OdfText(cited_text))
            new_nodes.append(mark_end)
            if after_text:
                new_nodes.append(_OdfText(after_text))

            # Slice-replace bypasses odfpy cache checks (safe for serialisation,
            # which walks childNodes directly)
            parent.childNodes[node_idx:node_idx + 1] = new_nodes
            for new_node in new_nodes:
                new_node.parentNode = parent

            return True

        offset = node_end

    return False


def _odt_para_in_body(para) -> bool:
    node = getattr(para, 'parentNode', None)
    while node is not None:
        if hasattr(node, 'qname') and node.qname in _ODT_NOTE_BODY_QNAMES:
            return False
        node = getattr(node, 'parentNode', None)
    return True


class OdtFieldWriter:
    """Injects Zotero reference marks into a LibreOffice .odt file."""

    def __init__(self, odt_path: str):
        if not _ODT_AVAILABLE:
            raise ImportError(
                "odfpy is required for .odt support. "
                "Run: python3 -m pip install odfpy"
            )
        self.odt_path = odt_path
        self._doc = _odf_load(odt_path)
        all_paras = list(self._doc.text.getElementsByType(_OdfP))
        self._paragraphs = [p for p in all_paras if _odt_para_in_body(p)]

    def apply_matches(
        self,
        results: List[MatchResult],
        accepted_suggestions: Optional[set] = None,
    ) -> None:
        accepted = accepted_suggestions or set()
        matched = [
            r for i, r in enumerate(results)
            if (r.matched or (r.is_suggestion and i in accepted))
            and not r.citation.in_footnote
        ]

        # Group compound citations sharing the same span
        groups = _group_by_span(matched)
        groups.sort(key=lambda g: (g[0].citation.paragraph_idx,
                                   -g[0].citation.char_start))

        for group in groups:
            cit = group[0].citation
            para_idx = cit.paragraph_idx
            if para_idx < 0 or para_idx >= len(self._paragraphs):
                continue
            para = self._paragraphs[para_idx]
            mark_name = self._build_mark_name(group, cit.raw_text)
            _inject_odt_marks(para, cit.char_start, cit.char_end, mark_name)

    def save(self, output_path: str) -> None:
        self._doc.save(output_path)

    def _build_mark_name(self, group: List[MatchResult], formatted_text: str) -> str:
        csl_citation = _build_csl_citation_json(group, formatted_text)
        return "ZOTERO_ITEM CSL_CITATION " + json.dumps(csl_citation)


# ---------------------------------------------------------------------------
# Convenience function
# ---------------------------------------------------------------------------

def write_zotero_document(
    source_path: str,
    output_path: str,
    results: List[MatchResult],
    accepted_suggestions: Optional[set] = None,
) -> int:
    """
    Creates output_path as a copy of source_path with matched citations
    (and any accepted suggestions) replaced by Zotero field codes.
    Supports both .docx and .odt source files.
    Returns the count of replacements made.
    """
    accepted = accepted_suggestions or set()
    effective_count = sum(
        1 for i, r in enumerate(results)
        if r.matched or (r.is_suggestion and i in accepted)
    )
    ext = Path(source_path).suffix.lower()
    if ext == '.odt':
        writer: OdtFieldWriter | FieldWriter = OdtFieldWriter(source_path)
    else:
        writer = FieldWriter(source_path)
    writer.apply_matches(results, accepted_suggestions=accepted)
    writer.save(output_path)
    return effective_count
