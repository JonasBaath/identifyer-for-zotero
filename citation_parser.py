"""
citation_parser.py
Extracts plain-text citations from a Word .docx or LibreOffice .odt file.
Supports author-year style (Smith, 2023) and numbered style [1].
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Optional ODT support (odfpy)
# ---------------------------------------------------------------------------
try:
    from odf.opendocument import load as _odf_load
    from odf.text import P as _OdfP
    from odf import teletype as _odf_teletype
    _ODT_AVAILABLE = True
except ImportError:
    _ODT_AVAILABLE = False

# QName for text:note (footnote/endnote container inside a paragraph)
_ODT_NOTE_QNAME = ('urn:oasis:names:tc:opendocument:xmlns:text:1.0', 'note')
# QNames for footnote/endnote body elements (to detect non-body paragraphs)
_ODT_NOTE_BODY_QNAMES = {
    ('urn:oasis:names:tc:opendocument:xmlns:text:1.0', 'footnote-body'),
    ('urn:oasis:names:tc:opendocument:xmlns:text:1.0', 'endnote-body'),
}


def _odt_para_in_body(para) -> bool:
    """Return True if para is a body paragraph (not inside a footnote/endnote)."""
    node = getattr(para, 'parentNode', None)
    while node is not None:
        if hasattr(node, 'qname') and node.qname in _ODT_NOTE_BODY_QNAMES:
            return False
        node = getattr(node, 'parentNode', None)
    return True


def _odt_para_text(para) -> str:
    """Extract plain text from an ODF paragraph, skipping note elements."""
    parts: List[str] = []

    def _walk(elem):
        for child in elem.childNodes:
            if child.nodeType == 3:          # TEXT_NODE
                parts.append(child.data)
            elif child.nodeType == 1:        # ELEMENT_NODE
                if hasattr(child, 'qname') and child.qname == _ODT_NOTE_QNAME:
                    continue                 # skip footnote content
                _walk(child)

    _walk(para)
    return ''.join(parts)


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

@dataclass
class Citation:
    raw_text: str               # exact text as found, e.g. "(Smith, 2023; Jones, 2024)"
    authors: List[str]          # extracted last names for THIS unit
    year: str                   # "2023"
    style: str                  # "author-year" | "numbered"
    ref_num: Optional[int]      # for numbered style
    paragraph_idx: int          # 0-based index into doc.paragraphs
    char_start: int             # char offset within paragraph full text
    char_end: int               # exclusive end offset
    in_footnote: bool = False   # True if found in a footnote/endnote
    display_text: str = ""      # per-unit display, e.g. "(Smith, 2023)" from a compound
    ref_text: str = ""          # full reference entry text (numbered style only)
    has_et_al: bool = False     # True if "et al." appears after the author name
    prefix: str = ""            # text before this citation in the paren block (e.g. "see ")
    suffix: str = ""            # text after this citation in the paren block (e.g. ", italics added")
    locator: str = ""           # page/chapter/section number, e.g. "88" or "10–11"
    locator_label: str = "page" # CSL label: "page", "chapter", "section"

    def display(self) -> str:
        if self.style == "numbered":
            return f"[{self.ref_num}]"
        return self.display_text or self.raw_text


# ---------------------------------------------------------------------------
# Regex patterns
# ---------------------------------------------------------------------------

# Last-name token: capitalised word (may include hyphens/apostrophes)
_NAME = r"[A-Z][A-Za-zÀ-ÿ'\-]+"
# Nobiliary particles (van, de, von, etc.) that may precede a capitalised surname
_PARTICLE_WORDS = (
    "van", "von", "de", "du", "da", "di", "del", "della", "der", "den",
    "het", "dos", "das", "do", "le", "la", "ten", "ter", "te",
)
_PARTICLE_PREFIX = r"(?:(?:" + "|".join(_PARTICLE_WORDS) + r")\s+)*"
_PARTICLES_SET = frozenset(_PARTICLE_WORDS)
# A full surname: optional particle(s) + capitalised name
_FULL_NAME = r"(?:" + _PARTICLE_PREFIX + _NAME + r")"
# A name unit: one or more consecutive full names
# Allows multi-word organisation names like "Swedish Food Agency"
# and names with particles like "van Giesen" or "de la Cruz"
_NAME_UNIT = r"(?:" + _FULL_NAME + r"(?:\s+" + _FULL_NAME + r")*)"
# Optional author initials like "J. ", "J.K. ", "J.-M. " that precede a surname
# for disambiguation.  Used inside UNIT_RE and INLINE_RE.
_INITIALS_OPT = r"(?:(?:[A-Z]\.[\s-]*)+)?"
# Year: 4-digit year (optionally + letter suffix), OR non-numeric tokens
# like "forthcoming", "in press", "accepted", "n.d."
_YEAR_NUM = r"\d{4}[a-z]?"
_YEAR_WORD = (
    r"(?:[Ff]orthcoming|[Ii]n\s+[Pp]ress|[Aa]ccepted|[Nn]\.?\s*[Dd]\.?|[Uu]npublished)"
)
_YEAR = r"(?:" + _YEAR_NUM + r"|" + _YEAR_WORD + r")"


def _normalize_year(year_str: str) -> str:
    """Normalize a year token for use as a lookup key.

    Returns '2020' for numeric years, or a canonical label for non-numeric
    tokens ('forthcoming', 'inpress', 'accepted', 'nd').
    """
    if re.match(r"\d{4}", year_str):
        return year_str[:4]
    lower = re.sub(r"[\s.]+", "", year_str.lower())
    if lower.startswith("forth"):
        return "forthcoming"
    if "press" in lower or lower.startswith("inp"):
        return "inpress"
    if lower.startswith("accept"):
        return "accepted"
    if lower in ("nd", "nd"):
        return "nd"
    if lower.startswith("unpub"):
        return "unpublished"
    return lower
# Optional page/chapter/section locator.
# Accepts APA style (, p. 45 / , pp. 45-47), Chicago style (: 45),
# short form (, 45), and chapter/section (, ch. 3 / , sec. 2).
# Bare numbers (no p./pp. prefix) are limited to 1-3 digits to avoid
# consuming follow-on 4-digit years like ", 2024" in compound citations.
_PAGE = (
    r"(?:"
    r"(?:,|\s*:)\s*(?:pp?\.\s*|ch\.\s*|sec\.\s*)\d{1,4}(?:\s*[–\-]\s*\d{1,4})?"  # with prefix: any digits
    r"|"
    r"(?:,|\s*:)\s*\d{1,3}(?!\d)(?:\s*[–\-]\s*\d{1,3}(?!\d))?"  # bare: max 3 digits, no partial 4-digit
    r")?"
)

# Abbreviation for "and others" in an author list — covers Latin "et al."
# (incl. "et. al." typo) and Swedish "m.fl."/"m. fl."/"mfl" (med flera).
_ET_AL_ALT = r"(?:et\.?\s+al\.?|m\.?\s*fl\.?)"

# Individual unit within a compound paren — uses _NAME_UNIT so multi-word
# org names like "(Swedish Food Agency, 2025)" are captured in full.
# The separator between authors and year accepts either a comma or plain
# whitespace so that "et al. 2017" (no comma) is handled alongside "et al., 2017".
UNIT_RE = re.compile(
    r"(?P<authors>" + _INITIALS_OPT + _NAME_UNIT + r"(?:\s*(?:,\s*&|,\s*\band\b|[,&]|\band\b)\s*" + _INITIALS_OPT + _NAME_UNIT + r")*(?:\s+" + _ET_AL_ALT + r")?)"
    r"(?:\s*,\s*|\s+)(?P<year>" + _YEAR + r")" + _PAGE,
    re.UNICODE | re.IGNORECASE,
)

# Detects any "et al." or "m.fl." variant in an author string
_ET_AL_RE = re.compile(r"\b" + _ET_AL_ALT, re.IGNORECASE)

# Follow-on bare year inside a compound citation, e.g. the "; 2018a" part of
# "(Aschemann-Witzel et al. 2017; 2018a)".  Author is inherited from the
# most recent preceding UNIT_RE match in the same parenthetical.
BARE_YEAR_RE = re.compile(r"[;,]\s*(?P<year>" + _YEAR + r")\b", re.IGNORECASE)

# Inline narrative: Smith (2023) — still anchored to a single _NAME word;
# multi-word org names are extended backwards in _parse_paragraph.
# Supports:  Smith (2023)         — basic
#             Smith's (2023)       — possessive (ASCII or curly apostrophe)
#             J. Smith (2023)      — with initials (optional)
#             Smith (2023, p. 45)  — with page/locator after year
INLINE_RE = re.compile(
    _INITIALS_OPT +
    r"(?P<author>" + _NAME + r")"
    r"(?:['\u2019]s)?"                          # optional possessive 's
    r"\s+\((?P<year>" + _YEAR + r")"
    r"(?:(?:[,:]\s*(?:pp?\.\s*|ch\.\s*|sec\.\s*)\d{1,4}(?:\s*[–\-]\s*\d{1,4})?)"  # prefixed locator
    r"|(?:[,:]\s*\d{1,3}(?!\d)(?:\s*[–\-]\s*\d{1,3}(?!\d))?))?"                  # bare locator (max 3 digits)
    r"\)",
    re.UNICODE,
)

# Capitalised words that should NOT be used to extend an author name backwards
# (sentence-starting articles / prepositions that look like proper nouns)
_INLINE_STOPWORDS = {
    "The", "A", "An", "As", "In", "On", "At", "By", "For", "To", "Of",
    "And", "Or", "But", "So", "If", "Yet", "Nor",
    "With", "From", "This", "That", "These", "Those",
    "While", "When", "Where", "Since", "After", "Before",
    "However", "Furthermore", "Moreover", "Although", "Because",
    "According", "Also", "Both", "Each", "Most", "Such",
    "Through", "During", "Between", "Within", "Without",
    "About", "Above", "Below", "Into", "Over", "Under",
}

# Numbered inline: [1] [1,2] [1-3] [1, 2, 3]
NUMBERED_INLINE_RE = re.compile(
    r"\[(?P<nums>\d+(?:\s*[,–\-]\s*\d+)*)\]"
)

# Reference list entry: starts with "1." or "[1]" optionally
REF_LIST_ENTRY_RE = re.compile(
    r"^\s*(?:\[(?P<num1>\d+)\]|(?P<num2>\d+)\.)\s+(?P<text>.+)$"
)


# Extracts the numeric locator (page/chapter/section) from the text that
# follows a year in a citation match, e.g. ":88", ", p. 45", ":10–11".
_LOCATOR_EXTRACT_RE = re.compile(
    r"(?:,|\s*:)\s*(?P<prefix>pp?\.\s*|ch\.\s*|sec\.\s*)?(?P<num>\d{1,4}(?:\s*[–\-]\s*\d{1,4})?)",
    re.UNICODE,
)


def _extract_locator(text_after_year: str):
    """Return (locator_value, csl_label) from text following the year in a match."""
    m = _LOCATOR_EXTRACT_RE.search(text_after_year)
    if not m:
        return "", "page"
    num = m.group("num").strip()
    prefix = (m.group("prefix") or "").lower().strip().rstrip(".")
    if prefix == "ch":
        label = "chapter"
    elif prefix == "sec":
        label = "section"
    else:
        label = "page"
    return num, label


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _split_authors(raw: str) -> List[str]:
    """Turn 'Smith & Jones' or 'Smith, Jones' or 'Smith et al.' / 'Smith m.fl.' into list of last names."""
    raw = re.sub(r"\s+" + _ET_AL_ALT, "", raw, flags=re.IGNORECASE)
    parts = re.split(r"\s*(?:[,&]|\band\b)\s*", raw)
    # Strip leading initials like "J." or "J.K." that some styles include
    # for disambiguation — we only need the surname for matching.
    cleaned = []
    for p in parts:
        p = re.sub(r'^(?:[A-Z]\.[\s-]*)+', '', p).strip()
        if p:
            cleaned.append(p)
    return cleaned


def _first_surname_key(author: str) -> str:
    """Extract the first capitalised surname from an author name, skipping particles."""
    # Split into words and skip leading particles (le, van, de, …)
    words = author.split()
    for w in words:
        if w.lower() not in _PARTICLES_SET:
            # Return the first non-particle word
            return w.lower()
    # Fallback: return first word if all words are particles (shouldn't happen)
    return words[0].lower() if words else author.lower()


def _build_display_text(authors: List[str], year: str, has_et_al: bool = False) -> str:
    """Reconstruct a minimal citation string like (Smith, 2023) from components."""
    if not authors:
        return f"({year})" if year else ""
    if has_et_al:
        # Preserve "et al." exactly as the source had it
        return f"({authors[0]} et al., {year})"
    if len(authors) == 1:
        return f"({authors[0]}, {year})"
    if len(authors) == 2:
        return f"({authors[0]} & {authors[1]}, {year})"
    return f"({authors[0]} et al., {year})"


# Regex that validates whether gap text is a genuine citation prefix
# (e.g., "see", "cf.", "e.g.,", "see also") vs. unmatched co-author names.
# Only sequences of known modifier words separated by punctuation/space are valid.
_VALID_PREFIX_RE = re.compile(
    r'^\s*(?:(?:see|also|e\.?g\.?|cf\.?|i\.?e\.?|compare|but|contra|pace|'
    r'hereafter|as|cited|in|following|after)\b[,;.\s]*)+$',
    re.IGNORECASE,
)


def _extract_prefix_suffix(
    inner_clean: str,
    all_positions: List[tuple],
) -> dict:
    """Compute per-citation prefix/suffix from gap text in a parenthetical.

    *all_positions* is a sorted list of (start, end, result_index, is_bare)
    tuples representing UNIT_RE and BARE_YEAR_RE matches within *inner_clean*.

    Returns a dict mapping result_index → (prefix_str, suffix_str).
    Bare-year entries participate in gap calculation but get empty prefix/suffix.
    Only text matching known prefix patterns (see, cf., e.g., etc.) is kept as
    prefix; other gap text (e.g. unmatched co-author names) is discarded.
    """
    all_positions = sorted(all_positions, key=lambda x: x[0])
    out: dict = {}

    for i, (start, end, idx, is_bare) in enumerate(all_positions):
        # --- prefix: gap between previous match end and this match start ---
        prev_end = all_positions[i - 1][1] if i > 0 else 0
        gap_before = inner_clean[prev_end:start]
        # Strip leading separators (;,) and whitespace
        prefix_candidate = re.sub(r'^[;,\s]+', '', gap_before)
        # Only accept known citation modifiers as prefix
        if prefix_candidate and _VALID_PREFIX_RE.match(prefix_candidate):
            prefix = prefix_candidate
        else:
            prefix = ""

        # --- suffix: gap between this match end and next match start ---
        next_start = all_positions[i + 1][0] if i + 1 < len(all_positions) else len(inner_clean)
        gap_after = inner_clean[end:next_start]
        # Take text up to the first ';' (the rest is prefix of the next citation)
        semi_pos = gap_after.find(';')
        suffix_raw = gap_after[:semi_pos] if semi_pos >= 0 else gap_after
        suffix = suffix_raw.strip()

        if is_bare:
            out[idx] = ("", "")
        else:
            out[idx] = (prefix, suffix)

    return out


def _get_paragraph_text(para) -> str:
    """Full plain text of a paragraph.

    - Includes text inside tracked-change insertion marks (<w:ins>).
    - Excludes text inside deletion marks (<w:del>).
    - Excludes display/result text of existing Zotero field codes so that
      already-coded citations are invisible to the parser and are not
      re-processed (which would corrupt the field code structure).

    Uses document-order iteration so all run wrappers (hyperlinks, ins, etc.)
    are handled correctly.
    """
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    del_tag     = f"{{{W}}}del"
    fldChar_tag = f"{{{W}}}fldChar"
    instr_tag   = f"{{{W}}}instrText"
    t_tag       = f"{{{W}}}t"

    parts: list = []
    in_zotero_result = False  # True after fldChar separate of a ZOTERO_ITEM field
    zotero_field     = False  # True once we've seen a ZOTERO_ITEM instrText

    for elem in para._element.iter():
        if elem.tag == fldChar_tag:
            ftype = elem.get(f"{{{W}}}fldCharType", "")
            if ftype == "begin":
                zotero_field = False
                in_zotero_result = False
            elif ftype == "separate":
                if zotero_field:
                    in_zotero_result = True
            elif ftype == "end":
                in_zotero_result = False
                zotero_field = False

        elif elem.tag == instr_tag:
            if elem.text and "ZOTERO_ITEM" in elem.text:
                zotero_field = True

        elif elem.tag == t_tag:
            if in_zotero_result or not elem.text:
                continue
            # Skip text inside <w:del>
            ancestor = elem.getparent()
            in_del = False
            while ancestor is not None and ancestor != para._element:
                if ancestor.tag == del_tag:
                    in_del = True
                    break
                ancestor = ancestor.getparent()
            if not in_del:
                parts.append(elem.text)

    return "".join(parts)


# ---------------------------------------------------------------------------
# Main parser
# ---------------------------------------------------------------------------

class CitationParser:
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self._doc: Optional[Document] = None
        self._para_texts: List[str] = []                 # normalised paragraph strings
        self._ref_start: Optional[int] = None            # paragraph index where references begin
        self.reference_map: dict[int, str] = {}          # num → reference text (numbered style)
        self.author_year_ref_map: dict[tuple, str] = {}  # (last_lower, year) → full entry text

    # Markers in instrText that identify citation-manager field codes
    _CITATION_FIELD_MARKERS = ("ZOTERO_ITEM", "EN.CITE", "EN.REFLIST")

    def _strip_citation_fields(self) -> None:
        """Replace Zotero/EndNote field codes with their visible text.

        This converts already-coded citations back to plain text so the
        parser can detect and re-match them against the Zotero library.
        Must be called after self._doc is loaded.
        """
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        w_r            = f"{{{W}}}r"
        w_fldChar      = f"{{{W}}}fldChar"
        w_instrText    = f"{{{W}}}instrText"
        w_fldCharType  = f"{{{W}}}fldCharType"
        w_p            = f"{{{W}}}p"

        # Process every <w:p> in the document body (including inside tables/SDTs)
        for p_elem in self._doc.element.body.iter(w_p):
            while True:
                children = list(p_elem)
                begin_idx = None
                is_citation_field = False
                sep_idx = None

                for i, child in enumerate(children):
                    if child.tag != w_r:
                        continue
                    fc = child.find(w_fldChar)
                    if fc is not None:
                        ftype = fc.get(w_fldCharType)
                        if ftype == "begin":
                            begin_idx = i
                            is_citation_field = False
                            sep_idx = None
                        elif ftype == "separate" and begin_idx is not None:
                            sep_idx = i
                        elif ftype == "end" and begin_idx is not None:
                            if is_citation_field and sep_idx is not None:
                                visible_runs = children[sep_idx + 1: i]
                                to_remove = children[begin_idx: i + 1]
                                insert_pos = begin_idx
                                for elem in to_remove:
                                    p_elem.remove(elem)
                                for j, vr in enumerate(visible_runs):
                                    p_elem.insert(insert_pos + j, vr)
                                break  # restart scan (indices changed)
                            begin_idx = None
                            sep_idx = None
                    else:
                        it = child.find(w_instrText)
                        if it is not None and begin_idx is not None:
                            if it.text and any(
                                m in it.text for m in self._CITATION_FIELD_MARKERS
                            ):
                                is_citation_field = True
                else:
                    break  # no more fields in this paragraph

    def parse(self) -> List[Citation]:
        ext = Path(self.docx_path).suffix.lower()
        if ext == '.odt':
            return self._parse_odt()

        # --- .docx path (original logic) ---
        self._doc = Document(self.docx_path)
        # Accept tracked changes first so that <w:ins>-wrapped runs become
        # direct children of <w:p>.  Without this, field codes inside tracked
        # insertions are invisible to _strip_citation_fields() (which only
        # scans direct children), causing char-offset mismatches with the
        # FieldWriter which also accepts changes before processing.
        from field_writer import accept_tracked_changes
        accept_tracked_changes(self._doc.element.body)
        # Strip existing Zotero/EndNote field codes so their visible citation
        # text becomes plain text that the parser can detect and re-match.
        self._strip_citation_fields()
        # Collect ALL paragraphs: body-level + inside table cells + inside SDTs
        self._all_paras = list(self._doc.paragraphs)
        # Include paragraphs from SDT (Structured Document Tag) content controls.
        # Zotero-generated bibliographies are often wrapped in an SDT block that
        # python-docx's .paragraphs property skips.
        from docx.text.paragraph import Paragraph as _Paragraph
        for sdt in self._doc.element.body.findall(qn('w:sdt')):
            sdt_content = sdt.find(qn('w:sdtContent'))
            if sdt_content is not None:
                for p_elem in sdt_content.findall(qn('w:p')):
                    self._all_paras.append(_Paragraph(p_elem, self._doc))
        for table in self._doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para not in self._all_paras:
                            self._all_paras.append(para)
        self._para_texts = [_get_paragraph_text(p) for p in self._all_paras]

        citations: List[Citation] = []

        # First pass: detect reference list and build number→text map
        self._ref_start = self._find_ref_section_start()
        self._build_reference_map()
        self._build_author_year_ref_map()

        # Second pass: extract citations from body paragraphs (skip reference section)
        body_end = self._ref_start if self._ref_start is not None else len(self._para_texts)
        for idx, text in enumerate(self._para_texts[:body_end]):
            citations.extend(self._parse_paragraph(text, idx, in_footnote=False))

        # Third pass: footnotes and endnotes
        citations.extend(self._parse_footnotes())

        # Deduplicate: keep only one Citation per (paragraph, char_start, display_text).
        # This prevents double-adding when both the compound-paren and single-unit
        # paths fire for the same match position.
        seen = set()
        unique: List[Citation] = []
        for c in citations:
            key = (c.paragraph_idx, c.char_start, c.display_text or c.raw_text)
            if key not in seen:
                seen.add(key)
                unique.append(c)

        # Populate ref_text for author-year citations from the bibliography map.
        # Use the first capitalised surname (skipping particles like "de", "van")
        # so that both org names and particle names map correctly.
        for c in unique:
            if c.style == "author-year" and not c.ref_text and c.authors:
                key = (_first_surname_key(c.authors[0]), _normalize_year(c.year))
                c.ref_text = self.author_year_ref_map.get(key, "")

        return unique

    # ------------------------------------------------------------------
    # Reference list detection
    # ------------------------------------------------------------------

    def _find_ref_section_start(self) -> int:
        """Locate the paragraph index where the reference/bibliography section begins.

        Strategy 1 — Validated heading (last wins):
            Scan for heading matches.  For each, verify that ≥3 AY_REF_RE entries
            exist within the next 25 paragraphs.  Keep the *last* validated heading
            so that an early false positive (e.g. a "Literature" section in the
            introduction) is ignored.

        Strategy 2 — Cluster scan (no heading needed):
            Scan forward from 50 % of the document looking for the longest cluster
            of AY_REF_RE-matching paragraphs, tolerating up to MAX_GAP consecutive
            non-matching paragraphs (blank lines, page breaks).  Require at least
            MIN_CLUSTER entries to avoid false positives from body text.
        """
        texts = self._para_texts
        n = len(texts)

        # --- Strategy 1: validated heading ---
        best_heading = None
        for i, raw in enumerate(texts):
            if raw.strip().lower() in self._REF_HEADINGS:
                count = sum(
                    1 for j in range(i + 1, min(n, i + 25))
                    if texts[j].strip() and self._AY_REF_RE.match(texts[j].strip())
                )
                if count >= 3:
                    best_heading = i
        if best_heading is not None:
            return best_heading

        # --- Strategy 2: cluster scan (handles missing heading, page breaks) ---
        MAX_GAP = 3      # tolerate up to 3 blank/non-matching paragraphs in a row
        MIN_CLUSTER = 5   # require at least 5 entries to count as a bibliography
        start_search = max(0, n // 2)

        best_start = n
        best_count = 0
        i = start_search
        while i < n:
            t = texts[i].strip()
            if t and self._AY_REF_RE.match(t):
                cluster_start = i
                cluster_count = 1
                gap = 0
                j = i + 1
                while j < n:
                    tj = texts[j].strip()
                    if tj and self._AY_REF_RE.match(tj):
                        cluster_count += 1
                        gap = 0
                    else:
                        gap += 1
                    if gap > MAX_GAP:
                        break
                    j += 1
                if cluster_count > best_count:
                    best_count = cluster_count
                    best_start = cluster_start
                i = j
            else:
                i += 1

        if best_count >= MIN_CLUSTER:
            return best_start

        # No reference section found — treat the entire document as body text.
        return n

    def _build_reference_map(self):
        """Scan the reference section to find numbered reference entries."""
        texts = self._para_texts
        # Start from _ref_start itself (headings won't match REF_LIST_ENTRY_RE)
        ref_start = self._ref_start if self._ref_start is not None else 0

        for raw in texts[ref_start:]:
            text = raw.strip()
            m = REF_LIST_ENTRY_RE.match(text)
            if m:
                entry_text = m.group("text").strip()
                # Require a 4-digit year to distinguish references from plain numbered lists
                if re.search(r'\b\d{4}\b', entry_text):
                    num = int(m.group("num1") or m.group("num2"))
                    self.reference_map[num] = entry_text

    # Matches an author-year bibliography entry, e.g.:
    #   Smith, J. (2023). Title of work. Journal, 10(2), 1–10.
    #   Smith, J., & Jones, A. (2023). ...
    _AY_REF_RE = re.compile(
        r"^" + _PARTICLE_PREFIX +                    # skip optional nobiliary particles
        r"(?P<last>[A-Z][A-Za-zÀ-ÿ'\-]{1,30})"     # first author capitalised surname
        r".*?(?:\(|\.\s*|\,\s*)"                     # year preceded by '(' or '. ' or ', '
        r"(?P<year>" + _YEAR + r")"                  # year (numeric or word)
        r"(?:\)|[.\s,;])",                           # year followed by ')' or separator
        re.UNICODE | re.IGNORECASE,
    )
    # Headings that mark the start of a bibliography section
    _REF_HEADINGS = {
        "references", "bibliography", "works cited", "literature cited",
        "reference list", "sources", "literature", "referanser",
        "litteraturliste", "referencias", "bibliographie",
    }

    def _build_author_year_ref_map(self):
        """Scan the bibliography/references section and map (first_author_lower, year) → full text."""
        texts = self._para_texts
        # Start from _ref_start itself (headings like "References" won't match _AY_REF_RE)
        ref_start = self._ref_start if self._ref_start is not None else 0

        for raw in texts[ref_start:]:
            text = raw.strip()
            if not text:
                continue
            m = self._AY_REF_RE.match(text)
            if m:
                last = m.group("last").lower()
                year = _normalize_year(m.group("year"))
                key = (last, year)
                if key not in self.author_year_ref_map:
                    self.author_year_ref_map[key] = text

    # ------------------------------------------------------------------
    # Per-paragraph parsing
    # ------------------------------------------------------------------

    def _parse_paragraph(
        self, text: str, para_idx: int, in_footnote: bool
    ) -> List[Citation]:
        results: List[Citation] = []

        # --- Author-year: parenthetical ---
        # Allow one level of nested parens so that
        # "(Intergovernmental Panel on Climate Change (IPCC) 2022)"
        # is captured alongside simple "(Smith 2020)" citations.
        _PAREN_RE = re.compile(r"\((?:[^()]*\([^()]*\))*[^()]*\)")
        for m in _PAREN_RE.finditer(text):
            chunk = m.group()
            if len(chunk) < 5 or len(chunk) > 502:
                continue
            chunk_start = m.start()
            if not re.search(r"\d{4}|[Ff]orthcoming|[Ii]n\s+[Pp]ress|[Aa]ccepted|[Nn]\.?\s*[Dd]\.?|[Uu]npublished", chunk):
                continue

            # Strip inner parenthetical abbreviations like "(IPCC)" so
            # UNIT_RE can parse "Intergovernmental Panel on Climate Change IPCC 2022"
            inner = chunk[1:-1]
            inner_clean = re.sub(r"\([^()]*\)", lambda mm: mm.group()[1:-1], inner)
            # Strip editor/translator markers (Ed., Eds., Trans.) that disrupt
            # the author-year pattern, e.g. "Smith, Ed., 2023" → "Smith 2023"
            inner_clean = re.sub(
                r',?\s*\b(?:Eds?|eds?|Trans|trans)\.\s*,?', '', inner_clean
            )
            sub_units = list(UNIT_RE.finditer(inner_clean))
            if not sub_units:
                continue

            # Collect citations and their positions for prefix/suffix extraction
            chunk_citations: List[Citation] = []
            all_positions: List[tuple] = []  # (start, end, idx, is_bare)

            for um in sub_units:
                raw_authors = um.group("authors")
                has_et_al = bool(_ET_AL_RE.search(raw_authors))
                authors = _split_authors(raw_authors)
                year = um.group("year")
                display = _build_display_text(authors, year, has_et_al)
                after_year = um.group(0)[um.end("year") - um.start():]
                loc_val, loc_label = _extract_locator(after_year)
                idx = len(chunk_citations)
                chunk_citations.append(Citation(
                    raw_text=chunk,
                    authors=authors,
                    year=year,
                    style="author-year",
                    ref_num=None,
                    paragraph_idx=para_idx,
                    char_start=chunk_start,
                    char_end=chunk_start + len(chunk),
                    in_footnote=in_footnote,
                    display_text=display,
                    has_et_al=has_et_al,
                    locator=loc_val,
                    locator_label=loc_label,
                ))
                all_positions.append((um.start(), um.end(), idx, False))

            # Handle bare follow-on years, e.g. "; 2018a" in
            # "(Aschemann-Witzel et al. 2017; 2018a)".
            # Carry the author forward from the most recent preceding unit.
            unit_spans = [(um.start(), um.end(), um) for um in sub_units]
            for bm in BARE_YEAR_RE.finditer(inner_clean):
                # Skip if this position falls inside an already-matched unit
                if any(s <= bm.start() < e for s, e, _ in unit_spans):
                    continue
                # Find the unit that starts closest before this bare year
                preceding = [(s, u) for s, e, u in unit_spans if s < bm.start()]
                if not preceding:
                    continue
                _, last_unit = max(preceding, key=lambda x: x[0])
                raw_authors = last_unit.group("authors")
                bare_authors = _split_authors(raw_authors)
                bare_has_et_al = bool(_ET_AL_RE.search(raw_authors))
                bare_year = bm.group("year")
                display = _build_display_text(bare_authors, bare_year, bare_has_et_al)
                idx = len(chunk_citations)
                chunk_citations.append(Citation(
                    raw_text=chunk,
                    authors=bare_authors,
                    year=bare_year,
                    style="author-year",
                    ref_num=None,
                    paragraph_idx=para_idx,
                    char_start=chunk_start,
                    char_end=chunk_start + len(chunk),
                    in_footnote=in_footnote,
                    display_text=display,
                    has_et_al=bare_has_et_al,
                ))
                all_positions.append((bm.start(), bm.end(), idx, True))

            # Extract prefix/suffix from gap text between matches
            if all_positions:
                ps_map = _extract_prefix_suffix(inner_clean, all_positions)
                for ci, (pfx, sfx) in ps_map.items():
                    chunk_citations[ci].prefix = pfx
                    chunk_citations[ci].suffix = sfx

            results.extend(chunk_citations)

        # --- Author-year: inline narrative (Smith (2023)) ---
        for m in INLINE_RE.finditer(text):
            author = m.group("author")
            # Strip possessive suffix: _NAME captures "Smith's" as one token
            # because ' is in the character class (needed for O'Brien etc.)
            author = re.sub(r"['\u2019]s$", "", author)
            year = m.group("year")
            name_start = m.start("author")

            # Extend the author name backwards through preceding capitalised
            # words so that "Swedish Food Agency (2025)" is captured in full
            # rather than just "Agency (2025)".
            _prev_word = re.compile(r"([A-Za-zÀ-ÿ'\-]+)\s*$")
            scan_pos = name_start
            while scan_pos > 0:
                before = text[:scan_pos]
                pm = _prev_word.search(before)
                if not pm:
                    break
                word = pm.group(1)
                if word[0].isupper() and word not in _INLINE_STOPWORDS:
                    author = word + " " + author
                    scan_pos = pm.start(1)
                    name_start = scan_pos
                elif word.lower() in _PARTICLES_SET:
                    author = word + " " + author
                    scan_pos = pm.start(1)
                    name_start = scan_pos
                else:
                    break

            display = _build_display_text([author], year)
            after_year = m.group(0)[m.end("year") - m.start():].rstrip(")")
            loc_val, loc_label = _extract_locator(after_year)
            results.append(Citation(
                raw_text=text[name_start:m.end()],
                authors=[author],
                year=year,
                style="author-year",
                ref_num=None,
                paragraph_idx=para_idx,
                char_start=name_start,
                char_end=m.end(),
                in_footnote=in_footnote,
                display_text=display,
                locator=loc_val,
                locator_label=loc_label,
            ))

        # --- Numbered ---
        for m in NUMBERED_INLINE_RE.finditer(text):
            nums_raw = m.group("nums")
            nums = self._expand_num_range(nums_raw)
            for n in nums:
                ref_text = self.reference_map.get(n, "")
                authors, year = self._parse_reference_entry(ref_text)
                results.append(Citation(
                    raw_text=m.group(),
                    authors=authors,
                    year=year,
                    style="numbered",
                    ref_num=n,
                    paragraph_idx=para_idx,
                    char_start=m.start(),
                    char_end=m.end(),
                    in_footnote=in_footnote,
                    display_text=f"[{n}]",
                    ref_text=ref_text,
                ))

        return results

    # ------------------------------------------------------------------
    # Footnotes
    # ------------------------------------------------------------------

    def _parse_footnotes(self) -> List[Citation]:
        """Parse citations from footnotes *and* endnotes via raw XML."""
        if self._doc is None:
            return []
        results: List[Citation] = []
        import zipfile
        from lxml import etree

        try:
            with zipfile.ZipFile(self.docx_path) as z:
                for part_name, tag_local in (
                    ('word/footnotes.xml', 'footnote'),
                    ('word/endnotes.xml', 'endnote'),
                ):
                    if part_name not in z.namelist():
                        continue
                    xml = z.read(part_name)
                    root = etree.fromstring(xml)
                    tag = qn(f'w:{tag_local}')
                    for idx, note in enumerate(root.findall(tag)):
                        note_type = note.get(qn('w:type'), '')
                        if note_type in ('separator', 'continuationSeparator'):
                            continue
                        for p_elem in note.findall(qn('w:p')):
                            texts = []
                            for r in p_elem.findall('.//' + qn('w:t')):
                                if r.text:
                                    texts.append(r.text)
                            text = ''.join(texts)
                            if text.strip():
                                results.extend(
                                    self._parse_paragraph(text, -(idx + 1), in_footnote=True)
                                )
        except (OSError, KeyError):
            pass
        return results

    # ------------------------------------------------------------------
    # ODT parsing
    # ------------------------------------------------------------------

    def _parse_odt(self) -> List[Citation]:
        """Parse citations from a LibreOffice .odt file using odfpy."""
        if not _ODT_AVAILABLE:
            raise ImportError(
                "odfpy is required for .odt support. "
                "Run: python3 -m pip install odfpy"
            )

        odt_doc = _odf_load(self.docx_path)

        # Collect body paragraphs only (exclude footnote/endnote paragraphs)
        all_paras = list(odt_doc.text.getElementsByType(_OdfP))
        body_paras = [p for p in all_paras if _odt_para_in_body(p)]

        self._para_texts = [_odt_para_text(p) for p in body_paras]

        self._ref_start = self._find_ref_section_start()
        self._build_reference_map()
        self._build_author_year_ref_map()

        # Extract citations from body paragraphs only (skip reference section)
        body_end = self._ref_start if self._ref_start is not None else len(self._para_texts)
        citations: List[Citation] = []
        for idx, text in enumerate(self._para_texts[:body_end]):
            citations.extend(self._parse_paragraph(text, idx, in_footnote=False))

        # Footnotes and endnotes in ODT
        citations.extend(self._parse_odt_notes(odt_doc))

        # Dedup (same logic as DOCX path)
        seen: set = set()
        unique: List[Citation] = []
        for c in citations:
            key = (c.paragraph_idx, c.char_start, c.display_text or c.raw_text)
            if key not in seen:
                seen.add(key)
                unique.append(c)

        # Populate ref_text from bibliography map
        for c in unique:
            if c.style == "author-year" and not c.ref_text and c.authors:
                key = (_first_surname_key(c.authors[0]), _normalize_year(c.year))
                c.ref_text = self.author_year_ref_map.get(key, "")

        return unique

    def _parse_odt_notes(self, odt_doc) -> List[Citation]:
        """Parse citations from footnotes and endnotes in an ODT document."""
        results: List[Citation] = []
        all_paras = list(odt_doc.text.getElementsByType(_OdfP))
        note_paras = [p for p in all_paras if not _odt_para_in_body(p)]
        for idx, para in enumerate(note_paras):
            text = _odf_teletype.extractText(para)
            if text and text.strip():
                results.extend(self._parse_paragraph(text, -(idx + 1), in_footnote=True))
        return results

    # ------------------------------------------------------------------
    # Utilities
    # ------------------------------------------------------------------

    def get_uncited_references(self, citations: List[Citation]) -> List[str]:
        """Return reference list entries that have no corresponding citation in the text."""
        uncited: List[str] = []

        # Numbered style
        if self.reference_map:
            cited_nums = {
                c.ref_num for c in citations
                if c.style == "numbered" and c.ref_num is not None
            }
            for num in sorted(self.reference_map):
                if num not in cited_nums:
                    uncited.append(f"[{num}] {self.reference_map[num]}")

        # Author-year style
        if self.author_year_ref_map:
            cited_keys: set = set()
            for c in citations:
                if c.style == "author-year" and c.authors:
                    cited_keys.add((_first_surname_key(c.authors[0]), _normalize_year(c.year)))
            for (last, year), text in self.author_year_ref_map.items():
                if (last, year) not in cited_keys:
                    uncited.append(text)

        return uncited

    @staticmethod
    def _expand_num_range(nums_raw: str) -> List[int]:
        """'1, 2, 3' or '1-3' or '1–3' → [1, 2, 3]"""
        nums_raw = nums_raw.strip()
        result: List[int] = []
        for part in re.split(r",\s*", nums_raw):
            range_m = re.match(r"(\d+)\s*[–\-]\s*(\d+)", part)
            if range_m:
                start, end = int(range_m.group(1)), int(range_m.group(2))
                result.extend(range(start, end + 1))
            else:
                try:
                    result.append(int(part.strip()))
                except ValueError:
                    pass
        return result

    @staticmethod
    def _parse_reference_entry(text: str):
        """Extract (authors, year) from a reference list entry string."""
        if not text:
            return [], ""
        year_m = re.search(r"\b(\d{4}[a-z]?)\b", text)
        year = year_m.group(1) if year_m else ""
        author_chunk = text[:year_m.start()].strip() if year_m else text[:50]
        author_chunk = re.sub(r"[\s,.()\[\]]+$", "", author_chunk)
        authors = _split_authors(author_chunk) if author_chunk else []
        return authors, year
